import os
import time
import json
import tempfile
import traceback
import requests
import subprocess
import shutil
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime, timedelta
from flask import Flask, request, jsonify, send_from_directory, Response, send_file, redirect, url_for
from flask_cors import CORS
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_mail import Mail, Message
from dotenv import load_dotenv
from openai import OpenAI
from werkzeug.utils import secure_filename
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from playwright.sync_api import sync_playwright
from looker_extractor import LookerStudioExtractor
from bs4 import BeautifulSoup
import cloudinary
import cloudinary.uploader
import cloudinary.api
from models import db, User, Session, ActivityLog, UserStats, log_activity

# Load environment variables
load_dotenv()

# Configuration
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
ELEVENLABS_API_KEY = os.getenv("ELEVENLABS_API_KEY", "")
TOPVIEW_API_KEY = os.getenv("TOPVIEW_API_KEY", "")
TOPVIEW_UID = os.getenv("TOPVIEW_UID", "")
MODEL_ID = os.getenv("MODEL_ID", "gpt-4o")
ASSISTANT_ID = os.getenv("ASSISTANT_ID", "")
PORT = int(os.getenv("PORT", "5173"))
DEBUG = os.getenv("DEBUG", "False").lower() == "true"
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
ALLOWED_EXTENSIONS = {'csv', 'xlsx', 'xls', 'pdf', 'txt', 'json'}

# Email Configuration (optional - for sending emails directly)
SMTP_SERVER = os.getenv("SMTP_SERVER", "smtp.gmail.com")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USERNAME = os.getenv("SMTP_USERNAME", "")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD", "")

# Cloudinary Configuration (for hosting media publicly)
CLOUDINARY_CLOUD_NAME = os.getenv("CLOUDINARY_CLOUD_NAME", "")
CLOUDINARY_API_KEY = os.getenv("CLOUDINARY_API_KEY", "")
CLOUDINARY_API_SECRET = os.getenv("CLOUDINARY_API_SECRET", "")

# Initialize Cloudinary if credentials are provided
if CLOUDINARY_CLOUD_NAME and CLOUDINARY_API_KEY and CLOUDINARY_API_SECRET:
    cloudinary.config(
        cloud_name=CLOUDINARY_CLOUD_NAME,
        api_key=CLOUDINARY_API_KEY,
        api_secret=CLOUDINARY_API_SECRET
    )
    print(f"✓ Cloudinary configured: {CLOUDINARY_CLOUD_NAME}")
else:
    print("⚠ Cloudinary not configured - add credentials to .env file")

# Initialize OpenAI client
client = OpenAI(api_key=OPENAI_API_KEY)

# Initialize Flask app
app = Flask(__name__, static_url_path="", static_folder="static")
CORS(app)  # Enable CORS for API calls
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

# Database & Authentication Configuration
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'dev-secret-key-change-in-production')
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL', 'sqlite:///supa_reports.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(seconds=int(os.getenv('PERMANENT_SESSION_LIFETIME', 86400)))

# Flask-Mail Configuration
app.config['MAIL_SERVER'] = os.getenv('MAIL_SERVER', 'smtp.gmail.com')
app.config['MAIL_PORT'] = int(os.getenv('MAIL_PORT', 587))
app.config['MAIL_USE_TLS'] = os.getenv('MAIL_USE_TLS', 'True').lower() == 'true'
app.config['MAIL_USERNAME'] = os.getenv('SMTP_USERNAME', '')
app.config['MAIL_PASSWORD'] = os.getenv('SMTP_PASSWORD', '')
app.config['MAIL_DEFAULT_SENDER'] = os.getenv('MAIL_DEFAULT_SENDER', 'supachatglobal@gmail.com')

# Initialize extensions
db.init_app(app)
mail = Mail(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# User loader for Flask-Login
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Update session activity on every request
@app.before_request
def update_session_activity():
    """Update the current user's session last_active timestamp"""
    if current_user.is_authenticated:
        try:
            # Get the user's active session
            active_session = current_user.get_active_session()
            if active_session:
                active_session.last_active = datetime.utcnow()
                db.session.commit()
        except Exception as e:
            # Don't let session update errors break the request
            db.session.rollback()
            print(f"Error updating session activity: {e}")

# Create database tables
with app.app_context():
    db.create_all()
    print("✓ Database initialized")

# Create logs directory
LOGS_DIR = os.path.join(os.path.dirname(__file__), 'logs')
os.makedirs(LOGS_DIR, exist_ok=True)
USAGE_LOG_FILE = os.path.join(LOGS_DIR, 'usage_log.json')

# Create uploads directory for profile pictures
UPLOADS_DIR = os.path.join(os.path.dirname(__file__), 'static', 'uploads', 'profiles')
os.makedirs(UPLOADS_DIR, exist_ok=True)

# ============================================================================
# SYSTEM PROMPT
# ============================================================================
SYSTEM_PROMPT = """You are **Azi**, the SupaChat Data Analyst Agent — a strategic marketing analyst designed to translate campaign performance into clear, actionable insights.

Your job is not to produce technical analytics.
Your reports are written for **marketing and brand teams** who care about meaning, momentum, and next steps — not scientific data tables.
Every report should feel like a conversation with a strategist who understands marketing impact.

------------------------------------------------------------
🎯 PURPOSE
------------------------------------------------------------
Transform campaign performance data into a concise, insightful, and actionable **Supa Report** that:
1. Diagnoses what worked, what didn't, and why.
2. Interprets data through a marketing and creative lens.
3. Benchmarks results against relevant industry best practice.
4. Delivers practical, prioritized recommendations for improvement.

------------------------------------------------------------
🧩 WORKFLOW
------------------------------------------------------------
1. **Briefing & Context**
   - Always confirm the briefing form first.
     • Brand name
     • Competitors + social handles
     • Market / country
     • Campaign objectives + KPIs
     • Reporting period
     • Uploaded dashboards or reports
     • Relevant research sources (social, web, YouTube, news)
     • Key questions or hypotheses

2. **Data Analysis**
   - Use `file_search` and `code_interpreter` to summarize metrics (impressions, reach, CTR, CPM, CPV, CPC, VCR, conversions, engagement rate, sentiment).
   - Identify clear trends and outliers, not just averages.
   - Always interpret *why* results look the way they do — not just *what* they are.

   ⚠️ **CRITICAL: NO DATA HALLUCINATION**
   - ONLY analyze data from two sources:
     1. Uploaded files (CSV, XLSX, PDF) - analyzed with file_search and code_interpreter
     2. Dashboard links - automatically scraped and provided to you as comprehensive text data

   **Dashboard Data Details:**
   - Dashboards are scraped using advanced browser automation that:
     * Navigates through ALL tabs/pages to capture complete data
     * Scrolls each page fully to load lazy-loaded content
     * Extracts structured data (tables, metrics, KPIs, filters)
     * Uses OCR (Optical Character Recognition) to extract text from images and visual elements
   - Dashboard data is provided in structured format with:
     * Tables: Full data tables with headers and rows
     * Metrics: Key performance indicators and scorecard values
     * Charts: Visual data representations
     * OCR Text: Text extracted from images, canvas elements, and visual content
     * Navigation: List of tabs/pages explored
   - **IMPORTANT**: Dashboard data is as reliable as uploaded files - use it with full confidence
   - The scraper scrolls to the bottom of each page and captures data from all tabs
   - OCR-extracted text may contain metrics, labels, and values from visual elements

   **If NO data is provided** (no file upload AND no dashboard data extracted):
     * State clearly: "No data was provided for analysis."
     * DO NOT make up metrics, numbers, or performance data
     * DO NOT assume or estimate campaign performance
     * Recommend: "Please upload a data file or provide a public dashboard link"
   - NEVER fabricate metrics like impressions, clicks, conversions, CTR, etc.
   - If you don't have actual data, SAY SO explicitly

3. **Contextual Intelligence**
   - Use `web_search` for market or category context and platform updates (≤90 days).
   - Reference known benchmarks (e.g., typical YouTube CPV, Meta CTR, TikTok CPM) where useful.
   - Cite findings naturally; focus on *insight relevance* rather than research detail.

4. **Synthesis & Reporting**
   - Structure reports into six key sections, each blending data highlights, strategic interpretation, and recommended actions.

------------------------------------------------------------
🧱 REPORT STRUCTURE
------------------------------------------------------------
# [Brand Name] SupaReport
### Campaign Performance Report with Actionable Insights
⸻

**1. Audience Insights — Turning Reach into Relevance**
Summarize audience performance: who engaged most, where, and how.
Use plain language, highlight patterns by age, geography, or device.

**What this means:**
Explain what the data reveals about the brand's audience behavior and mindset.
Focus on audience quality, intent, and opportunity.

**Actionable Insights:**
✅ Three to four short, practical ideas that help refine targeting, creative relevance, or localization.
⸻

**2. Media Channel Performance — Maximizing Efficiency**
Summarize cross-channel results (YouTube, TikTok, CTV, etc.) focusing on cost efficiency and audience attention.
Include 2–3 key metrics that demonstrate efficiency or impact.

**What this means:**
Interpret channel roles in the funnel — awareness, consideration, or conversion — and what their synergy achieved.

**Actionable Insights:**
✅ Three concise recommendations on budget optimization, channel sequencing, or creative adaptation.
⸻

**3. Creative Performance — Keeping the Story Fresh**
Describe how creative formats and messaging performed across placements.
Identify signs of fatigue or standout formats that captured attention.

**What this means:**
Highlight creative learnings: what resonated emotionally, what lost relevance, and how storytelling affected engagement.

**Actionable Insights:**
✅ Three creative strategies to sustain interest — rotation cadence, visual refresh, or storytelling evolution.
⸻

**4. Conversion & Engagement — From Attention to Action**
Summarize performance metrics like CTR, CVR, or engagement rates across devices and platforms.

**What this means:**
Translate engagement data into behavioral insights — who's converting, how, and under what creative or context conditions.

**Actionable Insights:**
✅ Three or more steps to improve engagement flow, optimize landing experience, or drive conversions.
⸻

**5. Optimization & Continuous Learning**
Summarize lessons learned from performance trends, creative testing, and channel mix.

**What this means:**
Discuss campaign maturity — where the strategy is gaining momentum and where refinements are needed.

**Actionable Insights:**
✅ Three to four forward-looking actions related to creative testing, targeting automation, analytics enhancement, or reporting cadence.
⸻

**6. Strategic Takeaway — Powering the Next Campaign**
Conclude with a forward-focused paragraph that connects insight to strategic direction.
Emphasize the big learning that can shape upcoming campaigns.

**Core Action Plan:**
1. Four short, high-level priorities guiding the next phase of marketing investment.
⸻

**In one line:**
End with a single, headline-style statement capturing the campaign's key strategic learning or success formula.

------------------------------------------------------------
💡 INSIGHT PRINCIPLES
------------------------------------------------------------
• Focus on meaning, not measurement — tell the story behind the data.
• Always answer: *Why did this happen?* and *What should the brand do next?*
• Support every point with clear evidence or directional data.
• Highlight "what to stop, start, and scale."
• Recommendations must be practical, brand-relevant, and time-bound (e.g., "rotate creative every 2–3 weeks" rather than "refresh creative").
• Use confident, marketing-savvy language: conversational, concise, and strategic.
• Avoid unnecessary analytics jargon — your reader is a marketer, not a data scientist.

------------------------------------------------------------
🏁 MISSION STATEMENT
------------------------------------------------------------
Azi's mission is to deliver **Supa Reports** that move from *data to insight to action*.
Each report should read like a strategic debrief from a senior marketing planner — insightful, grounded in real performance, and focused on helping the brand grow through smarter creative, smarter targeting, and smarter media investment.
"""

# ============================================================================
# OUTPUT SCHEMA
# ============================================================================
# Define section_block schema
SECTION_BLOCK_SCHEMA = {
    "type": "object",
    "required": [
        "key_findings",
        "supporting_data",
        "research_context",
        "implications",
        "actions"
    ],
    "properties": {
        "key_findings": {
            "type": "array",
            "items": {"type": "string"}
        },
        "supporting_data": {
            "type": "array",
            "items": {"type": "string"}
        },
        "research_context": {
            "type": "array",
            "items": {"type": "string"}
        },
        "implications": {
            "type": "array",
            "items": {"type": "string"}
        },
        "actions": {
            "type": "array",
            "items": {"type": "string"}
        }
    },
    "additionalProperties": False
}

OUTPUT_SCHEMA = {
    "type": "json_schema",
    "json_schema": {
        "name": "azi_analysis_output",
        "strict": True,
        "schema": {
            "type": "object",
            "required": [
                "audience",
                "media",
                "creative",
                "conversion",
                "competitive",
                "optimization",
                "bonus",
                "citations"
            ],
            "properties": {
                "audience": SECTION_BLOCK_SCHEMA,
                "media": SECTION_BLOCK_SCHEMA,
                "creative": SECTION_BLOCK_SCHEMA,
                "conversion": SECTION_BLOCK_SCHEMA,
                "competitive": SECTION_BLOCK_SCHEMA,
                "optimization": SECTION_BLOCK_SCHEMA,
                "bonus": {
                    "type": "object",
                    "required": ["one_sentence", "key_takeaway", "unexpected_learning"],
                    "properties": {
                        "one_sentence": {"type": "string"},
                        "key_takeaway": {"type": "string"},
                        "unexpected_learning": {"type": "string"}
                    },
                    "additionalProperties": False
                },
                "citations": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "required": ["title", "url", "source_type", "note"],
                        "properties": {
                            "title": {"type": "string"},
                            "url": {"type": "string"},
                            "source_type": {
                                "type": "string",
                                "enum": ["news", "website", "youtube", "social", "report"]
                            },
                            "note": {"type": "string"}
                        },
                        "additionalProperties": False
                    }
                }
            },
            "additionalProperties": False
        }
    }
}

# ============================================================================
# USAGE LOGGING
# ============================================================================

def log_usage(event_type, data):
    """
    Log usage events to JSON file for tracking and analytics.

    Args:
        event_type: Type of event (e.g., 'report_generation', 'video_generation')
        data: Dictionary with event-specific data
    """
    try:
        log_entry = {
            'timestamp': datetime.now().isoformat(),
            'event_type': event_type,
            **data
        }

        # Read existing logs
        logs = []
        if os.path.exists(USAGE_LOG_FILE):
            try:
                with open(USAGE_LOG_FILE, 'r') as f:
                    logs = json.load(f)
            except (json.JSONDecodeError, ValueError):
                logs = []

        # Append new log entry
        logs.append(log_entry)

        # Write back to file
        with open(USAGE_LOG_FILE, 'w') as f:
            json.dump(logs, f, indent=2)

        print(f"✓ Logged {event_type} event")

    except Exception as e:
        print(f"Error logging usage: {str(e)}")
        # Don't fail the request if logging fails


# ============================================================================
# PLAYWRIGHT DASHBOARD SCRAPING
# ============================================================================

def analyze_dashboards_with_playwright(dashboard_urls):
    """
    Use Playwright to navigate and extract insights from dashboard URLs.
    Works with Looker Studio dashboards and other web-based analytics platforms.
    Supports authenticated dashboards via persistent browser context.

    Args:
        dashboard_urls: List of dashboard URLs to analyze

    Returns:
        Dictionary with extracted data from each dashboard
    """
    if not dashboard_urls or len(dashboard_urls) == 0:
        return {}

    try:
        print(f"🌐 Scraping {len(dashboard_urls)} dashboard(s) with Playwright...")

        dashboard_insights = {}

        # Check for existing browser authentication data
        # Supports multiple browsers (firefox, chromium, webkit)
        base_dir = os.path.dirname(__file__)
        auth_dirs = {
            'firefox': os.path.join(base_dir, 'browser_data_firefox'),
            'chromium': os.path.join(base_dir, 'browser_data_chromium'),
            'webkit': os.path.join(base_dir, 'browser_data_webkit'),
        }

        # Find which browser has authentication data
        browser_type = None
        auth_dir = None
        for btype, bdir in auth_dirs.items():
            if os.path.exists(bdir) and os.listdir(bdir):
                browser_type = btype
                auth_dir = bdir
                print(f"  ℹ️ Found {btype.capitalize()} authentication data")
                break

        # If no auth data found, default to Chromium
        if not browser_type:
            browser_type = 'chromium'
            auth_dir = auth_dirs['chromium']
            print(f"  ℹ️ No authentication found, using {browser_type.capitalize()} (headless)")

        os.makedirs(auth_dir, exist_ok=True)

        with sync_playwright() as p:
            # Select the browser
            if browser_type == 'firefox':
                browser_engine = p.firefox
            elif browser_type == 'webkit':
                browser_engine = p.webkit
            else:  # chromium
                browser_engine = p.chromium

            # Browser arguments to bypass Google's automation detection
            browser_args = [
                '--disable-blink-features=AutomationControlled',
                '--disable-dev-shm-usage',
                '--no-sandbox',
                '--disable-setuid-sandbox',
                '--disable-web-security',
                '--disable-features=IsolateOrigins,site-per-process',
                '--disable-site-isolation-trials'
            ]

            # Use persistent context to maintain Google authentication
            # This allows the user to log in once and reuse the session
            # Can use headless mode since authentication is already saved
            headless_mode = os.getenv('SCRAPER_HEADLESS', 'true').lower() == 'true'

            try:
                context = browser_engine.launch_persistent_context(
                    auth_dir,
                    headless=headless_mode,  # Use headless by default (can be changed in .env)
                    viewport={'width': 1920, 'height': 1080},
                    user_agent='Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                    locale='en-US',
                    args=browser_args if browser_type == 'chromium' else []
                )

                print(f"  ℹ️ Using persistent {browser_type.capitalize()} context (headless={headless_mode})")

            except Exception as e:
                print(f"  ⚠️ Could not create persistent context: {e}")
                print(f"  ℹ️ Falling back to non-persistent mode")
                # Fallback to regular browser if persistent context fails
                browser = browser_engine.launch(
                    headless=True,
                    args=browser_args if browser_type == 'chromium' else []
                )
                context = browser.new_context(
                    viewport={'width': 1920, 'height': 1080},
                    user_agent='Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                    locale='en-US'
                )

            # Use existing page if available (persistent context may have one already)
            page = context.pages[0] if context.pages else context.new_page()

            # Remove webdriver property to hide automation
            page.add_init_script("""
                Object.defineProperty(navigator, 'webdriver', {
                    get: () => undefined
                });

                const originalQuery = window.navigator.permissions.query;
                window.navigator.permissions.query = (parameters) => (
                    parameters.name === 'notifications' ?
                        Promise.resolve({ state: Notification.permission }) :
                        originalQuery(parameters)
                );

                Object.defineProperty(navigator, 'plugins', {
                    get: () => [1, 2, 3, 4, 5]
                });

                Object.defineProperty(navigator, 'languages', {
                    get: () => ['en-US', 'en']
                });
            """)

            for idx, url in enumerate(dashboard_urls, 1):
                try:
                    print(f"  [{idx}/{len(dashboard_urls)}] Scraping: {url[:60]}...")

                    # Navigate to the dashboard
                    print(f"     Navigating to dashboard...")
                    # Use 'load' instead of 'networkidle' - Looker dashboards continuously fetch data
                    # and may never reach a true networkidle state
                    page.goto(url, wait_until='load', timeout=60000)  # 60 second timeout

                    # Add realistic delay to avoid detection
                    time.sleep(5)

                    # Check if we hit a login page or session expired
                    if 'accounts.google.com' in page.url:
                        print(f"  ❌ ERROR: Session has expired or is invalid.")
                        print(f"  ℹ️ Please run authentication setup again:")
                        print(f"      python3 setup_google_auth.py {browser_type}")
                        dashboard_insights[url] = "Error: Session expired. Please run setup_google_auth.py again."
                        continue

                    print(f"     Successfully loaded dashboard")
                    print(f"     Dashboard Title: {page.title()}")

                    # Create extractor instance
                    extractor = LookerStudioExtractor(page)

                    # Extract all data with navigation exploration and OCR enabled
                    dashboard_data = extractor.extract_all_data(
                        explore_nav=True,
                        enable_scrolling=True,
                        enable_ocr=True  # Enable OCR to extract data from canvas/images
                    )

                    # Format the extracted data as a text summary for OpenAI
                    summary_text = format_dashboard_data_as_text(dashboard_data)

                    dashboard_insights[url] = summary_text
                    print(f"  ✓ Successfully scraped dashboard {idx}")
                    print(f"     Extracted: {dashboard_data['summary']['total_tables']} tables, "
                          f"{dashboard_data['summary']['total_metrics']} metrics, "
                          f"{dashboard_data['summary']['total_charts']} charts")

                except Exception as e:
                    print(f"  ✗ Error scraping dashboard {idx}: {str(e)}")
                    dashboard_insights[url] = f"Error: Could not scrape dashboard - {str(e)}"

            context.close()

        print(f"✓ Completed dashboard scraping: {len(dashboard_insights)} processed")
        return dashboard_insights

    except Exception as e:
        print(f"Error in Playwright dashboard scraping: {str(e)}")
        traceback.print_exc()
        return {}


def fetch_url_content(url, timeout=10):
    """
    Fetch and extract text content from a URL.

    Args:
        url: The URL to fetch
        timeout: Request timeout in seconds

    Returns:
        Dictionary with URL, title, and extracted text
    """
    try:
        print(f"  Fetching URL: {url}")
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }

        response = requests.get(url, headers=headers, timeout=timeout)
        response.raise_for_status()

        # Parse HTML
        soup = BeautifulSoup(response.content, 'html.parser')

        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()

        # Get title
        title = soup.title.string if soup.title else url

        # Get text content
        text = soup.get_text(separator='\n', strip=True)

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        text = '\n'.join(line for line in lines if line)

        # Limit text length to avoid overwhelming context
        max_chars = 5000
        if len(text) > max_chars:
            text = text[:max_chars] + f"\n\n... (truncated, {len(text) - max_chars} more characters)"

        print(f"  ✓ Fetched {len(text)} characters from: {title}")

        return {
            'url': url,
            'title': title,
            'text': text,
            'success': True
        }

    except Exception as e:
        print(f"  ✗ Error fetching {url}: {str(e)}")
        return {
            'url': url,
            'error': str(e),
            'success': False
        }


def fetch_competitor_insights(competitor_urls):
    """
    Fetch content from competitor URLs and return formatted text.

    Args:
        competitor_urls: List of competitor URLs

    Returns:
        Formatted string with competitor insights
    """
    if not competitor_urls:
        return ""

    print(f"\n📊 Fetching competitor insights from {len(competitor_urls)} URLs...")

    results = []
    for url in competitor_urls[:5]:  # Limit to 5 URLs to avoid timeout
        result = fetch_url_content(url)
        results.append(result)

    # Format results
    lines = []
    lines.append("=" * 80)
    lines.append("COMPETITOR & RESEARCH INSIGHTS")
    lines.append("=" * 80)
    lines.append("")

    for idx, result in enumerate(results, 1):
        if result['success']:
            lines.append(f"[Source {idx}: {result['title']}]")
            lines.append(f"URL: {result['url']}")
            lines.append("-" * 80)
            lines.append(result['text'])
            lines.append("")
        else:
            lines.append(f"[Source {idx}: Failed to fetch]")
            lines.append(f"URL: {result['url']}")
            lines.append(f"Error: {result['error']}")
            lines.append("")

    return "\n".join(lines)


def format_dashboard_data_as_text(dashboard_data):
    """
    Convert dashboard data into text summary - PRIORITIZES OCR DATA.

    Args:
        dashboard_data: Dictionary with extracted dashboard data

    Returns:
        String with formatted text summary
    """
    lines = []

    # Add metadata
    metadata = dashboard_data.get('metadata', {})
    lines.append("=" * 80)
    lines.append("DASHBOARD DATA EXTRACTION")
    lines.append("=" * 80)
    if metadata.get('dashboard_title'):
        lines.append(f"Dashboard: {metadata['dashboard_title']}")
    lines.append(f"URL: {metadata.get('url', 'N/A')}")
    lines.append(f"Extracted: {metadata.get('timestamp', 'N/A')}")
    lines.append("")

    # Add navigation info first to show scope
    navigation = dashboard_data.get('navigation_explored', [])
    if navigation:
        lines.append(f"Tabs/Pages Explored: {len(navigation)}")
        lines.append(f"  → {', '.join(navigation)}")
        lines.append("")

    # PRIMARY DATA SOURCE: OCR extracted text - THIS IS THE MAIN DATA
    ocr_data = dashboard_data.get('ocr_text', [])
    if ocr_data:
        lines.append("=" * 80)
        lines.append("EXTRACTED DASHBOARD DATA (OCR from each tab)")
        lines.append("=" * 80)
        lines.append("")

        for idx, ocr_item in enumerate(ocr_data, 1):
            source = ocr_item.get('source', f'extraction_{idx}')
            text = ocr_item.get('text', '')
            char_count = ocr_item.get('char_count', 0)

            if text:
                lines.append(f"[Tab: {source}]")
                lines.append(f"Characters extracted: {char_count}")
                lines.append("-" * 80)
                lines.append(text)
                lines.append("")
                lines.append("=" * 80)
                lines.append("")
    else:
        lines.append("⚠️ WARNING: No OCR data extracted. Dashboard may be restricted or empty.")
        lines.append("")

    # Add summary at the end
    summary = dashboard_data.get('summary', {})
    lines.append("EXTRACTION SUMMARY")
    lines.append("-" * 80)
    if summary.get('total_ocr_extractions'):
        lines.append(f"✓ OCR Extractions: {summary.get('total_ocr_extractions', 0)} tabs")
        lines.append(f"✓ Total Characters: {summary.get('total_ocr_characters', 0):,}")
    lines.append(f"  Tables Found: {summary.get('total_tables', 0)}")
    lines.append(f"  Metrics Found: {summary.get('total_metrics', 0)}")
    lines.append(f"  Charts Found: {summary.get('total_charts', 0)}")

    return "\n".join(lines)


# ============================================================================
# Helper Functions
# ============================================================================

def allowed_file(filename):
    """Check if file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def validate_briefing(briefing):
    """Validate required briefing fields."""
    required_fields = ["brand", "market", "reporting_period", "objective"]
    errors = []

    for field in required_fields:
        if not briefing.get(field) or not briefing[field].strip():
            errors.append(f"Missing or empty required field: {field}")

    return errors


def ensure_assistant():
    """Create or retrieve the OpenAI Assistant."""
    global ASSISTANT_ID

    try:
        if ASSISTANT_ID:
            # Try to retrieve existing assistant
            try:
                assistant = client.beta.assistants.retrieve(ASSISTANT_ID)
                print(f"Using existing assistant: {ASSISTANT_ID}")
                return ASSISTANT_ID
            except Exception as e:
                print(f"Could not retrieve assistant {ASSISTANT_ID}: {e}")
                print("Creating new assistant...")

        # Create new assistant
        assistant = client.beta.assistants.create(
            name="SupaReports — Upgrade Your Insights",
            model=MODEL_ID,
            instructions=SYSTEM_PROMPT,
            tools=[
                {"type": "file_search"},
                {"type": "code_interpreter"}
            ],
            response_format=OUTPUT_SCHEMA
        )
        ASSISTANT_ID = assistant.id
        print(f"Created new assistant: {ASSISTANT_ID}")
        print(f"Add this to your .env file to reuse: ASSISTANT_ID={ASSISTANT_ID}")
        return ASSISTANT_ID

    except Exception as e:
        raise Exception(f"Failed to create/retrieve assistant: {str(e)}")


def poll_run_status(thread_id, run_id, timeout=300, poll_interval=1.5):
    """
    Poll the run status until completion or timeout.

    Args:
        thread_id: Thread ID
        run_id: Run ID
        timeout: Maximum time to wait in seconds (default 5 minutes)
        poll_interval: Time between status checks in seconds

    Returns:
        Final run object
    """
    start_time = time.time()

    while True:
        elapsed = time.time() - start_time
        if elapsed > timeout:
            raise TimeoutError(f"Analysis timed out after {timeout} seconds")

        run = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run_id)

        if run.status == "completed":
            return run
        elif run.status in ("failed", "cancelled", "expired"):
            error_msg = f"Run ended with status: {run.status}"
            if hasattr(run, 'last_error') and run.last_error:
                error_msg += f" - {run.last_error}"
            raise Exception(error_msg)

        # Continue polling
        time.sleep(poll_interval)


# ============================================================================
# Routes
# ============================================================================

# ============================================================================
# Authentication Routes
# ============================================================================

@app.route("/api/auth/signup", methods=["POST"])
def signup():
    """User registration endpoint"""
    try:
        data = request.get_json()
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')
        username = data.get('username', '').strip()

        # Validation
        if not email or not password:
            return jsonify({"error": "Email and password are required"}), 400

        if len(password) < 8:
            return jsonify({"error": "Password must be at least 8 characters"}), 400

        # Check if email already exists
        if User.query.filter_by(email=email).first():
            return jsonify({"error": "Email already registered"}), 400

        # Check if username already exists (if provided)
        if username and User.query.filter_by(username=username).first():
            return jsonify({"error": "Username already taken"}), 400

        # Create new user
        user = User(
            email=email,
            username=username or email.split('@')[0],
            verified=False
        )
        user.set_password(password)
        verification_token = user.generate_verification_token()

        db.session.add(user)
        db.session.commit()

        # Send verification email
        try:
            verification_url = f"{request.host_url}api/auth/verify/{verification_token}"
            msg = Message(
                subject="Verify Your Supa Reports Account",
                recipients=[email],
                html=f"""
                <html>
                <body style="font-family: Arial, sans-serif; padding: 20px; background-color: #f4f4f4;">
                    <div style="max-width: 600px; margin: 0 auto; background-color: white; padding: 30px; border-radius: 10px;">
                        <h2 style="color: #333;">Welcome to Supa Reports!</h2>
                        <p>Thank you for signing up. Please verify your email address by clicking the button below:</p>
                        <p style="text-align: center; margin: 30px 0;">
                            <a href="{verification_url}"
                               style="background-color: #4CAF50; color: white; padding: 12px 30px;
                                      text-decoration: none; border-radius: 5px; display: inline-block;">
                                Verify Email
                            </a>
                        </p>
                        <p style="color: #666; font-size: 12px;">
                            If the button doesn't work, copy and paste this link into your browser:<br>
                            <a href="{verification_url}">{verification_url}</a>
                        </p>
                        <p style="color: #666; font-size: 12px;">
                            This verification link will expire in 24 hours.
                        </p>
                    </div>
                </body>
                </html>
                """
            )
            mail.send(msg)
        except Exception as e:
            print(f"Error sending verification email: {e}")
            # Don't fail signup if email fails
            pass

        return jsonify({
            "success": True,
            "message": "Account created! Please check your email to verify your account.",
            "email": email
        }), 201

    except Exception as e:
        db.session.rollback()
        print(f"Signup error: {e}")
        return jsonify({"error": "An error occurred during signup"}), 500


@app.route("/api/auth/verify/<token>", methods=["GET"])
def verify_email(token):
    """Email verification endpoint"""
    try:
        user = User.query.filter_by(verification_token=token).first()

        if not user:
            return """
            <html>
            <body style="font-family: Arial; text-align: center; padding: 50px;">
                <h2 style="color: #f44336;">Invalid Verification Link</h2>
                <p>This verification link is invalid or has already been used.</p>
            </body>
            </html>
            """, 400

        if user.verified:
            return """
            <html>
            <body style="font-family: Arial; text-align: center; padding: 50px;">
                <h2 style="color: #4CAF50;">Already Verified</h2>
                <p>This email address has already been verified.</p>
                <p><a href="/">Go to Login</a></p>
            </body>
            </html>
            """

        # Mark user as verified
        user.verified = True
        user.verification_token = None  # Clear token after use
        db.session.commit()

        return """
        <html>
        <body style="font-family: Arial; text-align: center; padding: 50px;">
            <h2 style="color: #4CAF50;">✓ Email Verified!</h2>
            <p>Your email has been successfully verified. You can now log in.</p>
            <p><a href="/" style="background-color: #4CAF50; color: white; padding: 10px 20px;
                   text-decoration: none; border-radius: 5px; display: inline-block; margin-top: 20px;">
                Go to Login
            </a></p>
        </body>
        </html>
        """

    except Exception as e:
        print(f"Verification error: {e}")
        return """
        <html>
        <body style="font-family: Arial; text-align: center; padding: 50px;">
            <h2 style="color: #f44336;">Verification Failed</h2>
            <p>An error occurred during verification. Please try again.</p>
        </body>
        </html>
        """, 500


@app.route("/api/auth/login", methods=["POST"])
def login():
    """User login endpoint with one-session-per-user enforcement"""
    try:
        data = request.get_json()
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')

        if not email or not password:
            return jsonify({"error": "Email and password are required"}), 400

        # Find user
        user = User.query.filter_by(email=email).first()

        if not user or not user.check_password(password):
            return jsonify({"error": "Invalid email or password"}), 401

        # Check if email is verified
        if not user.verified:
            return jsonify({
                "error": "Email not verified",
                "message": "Please check your email and verify your account before logging in."
            }), 403

        # Check for existing active session and deactivate it
        # This allows the user to login even if logout didn't complete properly
        existing_session = user.get_active_session()
        if existing_session:
            print(f"Deactivating existing session for user {user.email}")
            existing_session.deactivate()

        # Create new session
        session_token = Session.generate_token()
        new_session = Session(
            user_id=user.id,
            session_token=session_token,
            ip_address=request.remote_addr,
            user_agent=request.headers.get('User-Agent', ''),
            is_active=True
        )
        db.session.add(new_session)

        # Update last login
        user.last_login = datetime.utcnow()
        db.session.commit()

        # Log in user with Flask-Login
        login_user(user, remember=True)

        # Create user stats if doesn't exist
        if not user.stats:
            stats = UserStats(user_id=user.id)
            db.session.add(stats)
            db.session.commit()

        return jsonify({
            "success": True,
            "message": "Login successful",
            "user": {
                "id": user.id,
                "email": user.email,
                "username": user.username,
                "profile_picture": user.profile_picture
            },
            "session_token": session_token
        }), 200

    except Exception as e:
        db.session.rollback()
        print(f"Login error: {e}")
        return jsonify({"error": "An error occurred during login"}), 500


@app.route("/api/auth/logout", methods=["POST"])
@login_required
def logout():
    """User logout endpoint"""
    try:
        # Deactivate current session
        active_session = current_user.get_active_session()
        if active_session:
            active_session.deactivate()

        # Logout with Flask-Login
        logout_user()

        return jsonify({
            "success": True,
            "message": "Logged out successfully"
        }), 200

    except Exception as e:
        print(f"Logout error: {e}")
        return jsonify({"error": "An error occurred during logout"}), 500


@app.route("/api/auth/request-reset", methods=["POST"])
def request_password_reset():
    """Request password reset - sends reset email"""
    try:
        data = request.get_json()
        email = data.get('email', '').strip().lower()

        if not email:
            return jsonify({"error": "Email is required"}), 400

        # Find user by email
        user = User.query.filter_by(email=email).first()

        # Always return success to prevent email enumeration
        # Don't reveal whether email exists in database
        if user:
            # Generate reset token (expires in 24 hours)
            reset_token = user.generate_reset_token(expires_in_hours=24)
            db.session.commit()

            # Create reset link
            reset_url = f"http://localhost:5173?reset_token={reset_token}"

            # Send password reset email
            try:
                msg = Message(
                    subject="Reset Your Password - Supa Reports",
                    recipients=[email],
                    html=f"""
                    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                        <h2 style="color: #4CAF50;">Password Reset Request</h2>
                        <p>You requested to reset your password. Click the button below to set a new password:</p>
                        <p style="margin: 30px 0;">
                            <a href="{reset_url}"
                               style="background-color: #4CAF50; color: white; padding: 12px 24px;
                                      text-decoration: none; border-radius: 4px; display: inline-block;">
                                Reset Password
                            </a>
                        </p>
                        <p style="color: #666;">This link will expire in 24 hours.</p>
                        <p style="color: #666;">If you didn't request this, please ignore this email.</p>
                        <hr style="margin: 30px 0; border: none; border-top: 1px solid #ddd;">
                        <p style="color: #999; font-size: 12px;">
                            If the button doesn't work, copy and paste this link:<br>
                            <a href="{reset_url}">{reset_url}</a>
                        </p>
                    </div>
                    """
                )
                mail.send(msg)
                print(f"Password reset email sent to {email}")
            except Exception as e:
                print(f"Error sending reset email: {e}")
                # Continue anyway - don't reveal if email failed

        return jsonify({
            "success": True,
            "message": "If that email exists, we've sent password reset instructions."
        }), 200

    except Exception as e:
        print(f"Request reset error: {e}")
        traceback.print_exc()
        return jsonify({"error": "An error occurred processing your request"}), 500


@app.route("/api/auth/reset-password", methods=["POST"])
def reset_password():
    """Reset password using token"""
    try:
        data = request.get_json()
        token = data.get('token', '').strip()
        new_password = data.get('password', '')

        if not token or not new_password:
            return jsonify({"error": "Token and password are required"}), 400

        if len(new_password) < 8:
            return jsonify({"error": "Password must be at least 8 characters"}), 400

        # Find user with this reset token
        user = User.query.filter_by(reset_token=token).first()

        if not user:
            return jsonify({"error": "Invalid or expired reset token"}), 400

        # Verify token is valid and not expired
        if not user.verify_reset_token(token):
            return jsonify({"error": "Invalid or expired reset token"}), 400

        # Update password
        user.set_password(new_password)
        user.clear_reset_token()
        db.session.commit()

        print(f"Password reset successful for user: {user.email}")

        return jsonify({
            "success": True,
            "message": "Password reset successful! You can now log in with your new password."
        }), 200

    except Exception as e:
        db.session.rollback()
        print(f"Reset password error: {e}")
        traceback.print_exc()
        return jsonify({"error": "An error occurred resetting your password"}), 500


@app.route("/api/auth/online-users", methods=["GET"])
@login_required
def get_online_users():
    """Get all currently online users"""
    try:
        # Query all active sessions with their users
        # Active sessions are those within the last 15 minutes
        active_threshold = datetime.utcnow() - timedelta(minutes=15)

        active_sessions = Session.query.filter(
            Session.is_active == True,
            Session.last_active >= active_threshold
        ).all()

        # Build list of online users with their info
        online_users = []
        seen_user_ids = set()

        for session in active_sessions:
            if session.user_id not in seen_user_ids:
                user = session.user
                if user:
                    online_users.append({
                        'id': user.id,
                        'username': user.username,
                        'email': user.email,
                        'profile_picture': user.profile_picture,
                        'initials': user.username[0:2].upper() if user.username else user.email[0:2].upper(),
                        'last_active': session.last_active.isoformat()
                    })
                    seen_user_ids.add(session.user_id)

        # Sort by last active (most recent first)
        online_users.sort(key=lambda x: x['last_active'], reverse=True)

        return jsonify({
            'success': True,
            'count': len(online_users),
            'users': online_users
        }), 200

    except Exception as e:
        print(f"Error fetching online users: {e}")
        traceback.print_exc()
        return jsonify({"error": "Failed to fetch online users"}), 500


@app.route("/api/auth/session", methods=["GET"])
@login_required
def get_session_info():
    """Get current user session information"""
    try:
        user_stats = current_user.stats or UserStats(user_id=current_user.id)

        # Get total live sessions count
        total_sessions = Session.query.filter_by(is_active=True).count()

        # For admins, show total reports across all users
        # For regular users, show only their own reports
        if current_user.is_admin:
            # Sum up all reports from all users
            total_reports = db.session.query(
                db.func.sum(UserStats.reports_count)
            ).scalar() or 0
            reports_count = total_reports
        else:
            # Show only this user's reports
            reports_count = user_stats.reports_count

        return jsonify({
            "user": {
                "id": current_user.id,
                "email": current_user.email,
                "username": current_user.username,
                "profile_picture": current_user.profile_picture,
                "verified": current_user.verified,
                "is_admin": current_user.is_admin,
                "last_login": current_user.last_login.isoformat() if current_user.last_login else None
            },
            "stats": {
                "reports": reports_count,
                "audio": user_stats.audio_count,
                "videos": user_stats.video_count,
                "emails": user_stats.emails_sent_count,
                "analyses": user_stats.analyses_count
            },
            "global_stats": {
                "live_sessions": total_sessions
            }
        }), 200

    except Exception as e:
        print(f"Session info error: {e}")
        traceback.print_exc()
        return jsonify({"error": "Failed to get session info"}), 500


@app.route("/api/user/profile-picture", methods=["POST"])
@login_required
def upload_profile_picture():
    """Upload or select profile picture"""
    try:
        # Check if it's a file upload or default avatar selection
        if 'picture' in request.files:
            file = request.files['picture']

            if file.filename == '':
                return jsonify({"error": "No file selected"}), 400

            # Validate file type
            allowed_extensions = {'png', 'jpg', 'jpeg', 'gif'}
            ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''

            if ext not in allowed_extensions:
                return jsonify({"error": "Invalid file type. Only PNG, JPG, JPEG, GIF allowed"}), 400

            # Save file
            filename = f"user_{current_user.id}.{ext}"
            filepath = os.path.join(UPLOADS_DIR, filename)
            file.save(filepath)

            # Update user profile picture path
            current_user.profile_picture = f"/static/uploads/profiles/{filename}"
            db.session.commit()

            return jsonify({
                "success": True,
                "profile_picture": current_user.profile_picture
            }), 200

        elif 'avatar' in request.json:
            # Select default avatar
            avatar_num = request.json.get('avatar')
            if avatar_num not in [1, 2, 3, 4, 5]:
                return jsonify({"error": "Invalid avatar number"}), 400

            current_user.profile_picture = f"/static/avatars/default-{avatar_num}.png"
            db.session.commit()

            return jsonify({
                "success": True,
                "profile_picture": current_user.profile_picture
            }), 200

        else:
            return jsonify({"error": "No picture or avatar selection provided"}), 400

    except Exception as e:
        db.session.rollback()
        print(f"Profile picture upload error: {e}")
        return jsonify({"error": "Failed to upload profile picture"}), 500


# ============================================================================
# Main Routes
# ============================================================================

@app.route("/")
def index():
    """Serve the main HTML page."""
    return send_from_directory("static", "index.html")


@app.route("/assets/<path:filename>")
def serve_assets(filename):
    """Serve files from the assets directory."""
    assets_dir = os.path.join(os.path.dirname(__file__), 'assets')
    return send_from_directory(assets_dir, filename)


@app.route("/api/health", methods=["GET"])
def health_check():
    """Health check endpoint."""
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "model": MODEL_ID,
        "assistant_configured": bool(ASSISTANT_ID)
    })


@app.route("/api/send-email", methods=["POST"])
def send_email():
    """
    Send HTML email to recipients.
    Accepts JSON with: from_email, to_emails (list), subject, html_content
    """
    try:
        data = request.get_json()

        from_email = data.get('from_email')
        to_emails = data.get('to_emails', [])
        subject = data.get('subject')
        html_content = data.get('html_content')

        # Validation
        if not from_email:
            return jsonify({"error": "from_email is required"}), 400

        if not to_emails or len(to_emails) == 0:
            return jsonify({"error": "to_emails is required"}), 400

        if not subject:
            return jsonify({"error": "subject is required"}), 400

        if not html_content:
            return jsonify({"error": "html_content is required"}), 400

        # If SMTP credentials are configured, use them
        # Otherwise, use the from_email as sender (requires app-specific password)
        smtp_user = SMTP_USERNAME if SMTP_USERNAME else from_email
        smtp_pass = SMTP_PASSWORD

        if not smtp_pass:
            return jsonify({
                "error": "Email credentials not configured",
                "message": "Please configure SMTP_PASSWORD in .env file or use your email provider's app-specific password"
            }), 400

        # Create message
        msg = MIMEMultipart('alternative')
        msg['From'] = from_email
        msg['To'] = ', '.join(to_emails)
        msg['Subject'] = subject

        # Attach HTML content
        html_part = MIMEText(html_content, 'html')
        msg.attach(html_part)

        # Send email
        print(f"Sending email from {from_email} to {to_emails}...")

        # Try port 587 with STARTTLS first, fallback to port 465 with SSL
        try:
            with smtplib.SMTP(SMTP_SERVER, SMTP_PORT, timeout=10) as server:
                server.starttls()
                server.login(smtp_user, smtp_pass)
                server.send_message(msg)
        except (smtplib.SMTPException, TimeoutError, OSError) as smtp_error:
            print(f"Port {SMTP_PORT} failed, trying SSL on port 465...")
            # Fallback to SSL on port 465
            with smtplib.SMTP_SSL(SMTP_SERVER, 465, timeout=10) as server:
                server.login(smtp_user, smtp_pass)
                server.send_message(msg)

        print(f"✓ Email sent successfully to {len(to_emails)} recipient(s)")

        return jsonify({
            "success": True,
            "message": f"Email sent to {len(to_emails)} recipient(s)",
            "recipients": to_emails
        })

    except smtplib.SMTPAuthenticationError:
        return jsonify({
            "error": "Authentication failed",
            "message": "Invalid email credentials. For Gmail, use an app-specific password: https://support.google.com/accounts/answer/185833"
        }), 401

    except smtplib.SMTPException as e:
        return jsonify({
            "error": "SMTP error",
            "message": str(e)
        }), 500

    except Exception as e:
        print(f"Error sending email: {str(e)}")
        traceback.print_exc()
        return jsonify({
            "error": "Failed to send email",
            "message": str(e)
        }), 500


@app.route("/api/upload-to-cloudinary", methods=["POST"])
def upload_to_cloudinary():
    """
    Upload video to Cloudinary and optionally convert to GIF.
    Accepts video file and returns public URLs for both video and GIF.
    """
    try:
        if not CLOUDINARY_CLOUD_NAME:
            return jsonify({
                "error": "Cloudinary not configured",
                "message": "Please add Cloudinary credentials to .env file"
            }), 400

        # Check if video file is in the request
        if 'video' not in request.files:
            return jsonify({"error": "No video file provided"}), 400

        video_file = request.files['video']

        if video_file.filename == '':
            return jsonify({"error": "Empty filename"}), 400

        # Get optional parameters
        convert_to_gif = request.form.get('convert_to_gif', 'true').lower() == 'true'
        gif_duration = int(request.form.get('gif_duration', '5'))  # Default 5 seconds

        print(f"📤 Uploading video to Cloudinary: {video_file.filename}")

        # Save video temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as temp_video:
            video_file.save(temp_video.name)
            temp_video_path = temp_video.name

        try:
            # Upload original video to Cloudinary
            print("Uploading original video...")
            video_upload = cloudinary.uploader.upload(
                temp_video_path,
                resource_type="video",
                folder="supa_reports/videos",
                overwrite=True,
                notification_url=None
            )

            video_url = video_upload['secure_url']
            print(f"✓ Video uploaded: {video_url}")

            gif_url = None
            if convert_to_gif:
                print(f"Converting first {gif_duration} seconds to GIF...")

                # Use Cloudinary's transformation to create GIF
                # Extract first N seconds, optimize for email
                gif_public_id = video_upload['public_id']

                # Build GIF URL with transformations
                gif_url, _ = cloudinary.utils.cloudinary_url(
                    gif_public_id,
                    resource_type='video',
                    format='gif',
                    transformation=[
                        {'duration': gif_duration},         # First N seconds
                        {'width': 600, 'crop': 'scale'},    # Resize to 600px width
                        {'quality': 'auto:low'},            # Optimize file size
                        {'flags': 'animated'},              # Ensure it's animated
                        {'effect': 'loop'}                  # Infinite looping
                    ]
                )

                print(f"✓ GIF URL generated: {gif_url}")

            return jsonify({
                "success": True,
                "video_url": video_url,
                "gif_url": gif_url,
                "public_id": video_upload['public_id'],
                "duration": video_upload.get('duration', 0)
            })

        finally:
            # Clean up temporary file
            if os.path.exists(temp_video_path):
                os.unlink(temp_video_path)

    except Exception as e:
        print(f"Error uploading to Cloudinary: {str(e)}")
        traceback.print_exc()
        return jsonify({
            "error": "Upload failed",
            "message": str(e)
        }), 500


@app.route("/api/analyze", methods=["POST"])
def analyze():
    """
    Main analysis endpoint.
    Accepts briefing form data and optional file upload.
    Returns structured analysis JSON.
    """
    start_time = time.time()
    try:
        # Parse briefing fields from form data
        data = request.form.to_dict()

        # Parse comma-separated fields
        competitors_list = [x.strip() for x in data.get("competitors", "").split(",") if x.strip()]
        competitor_urls_list = [x.strip() for x in data.get("competitor_urls", "").split(",") if x.strip()]

        # Combine competitors and URLs into structured format
        competitors = []
        for i, name in enumerate(competitors_list):
            competitor_obj = {"name": name}
            if i < len(competitor_urls_list):
                competitor_obj["handles"] = [competitor_urls_list[i]]
            competitors.append(competitor_obj)

        dashboard_links = [x.strip() for x in data.get("dashboard_links", "").split(",") if x.strip()]
        research_urls = [x.strip() for x in data.get("research_urls", "").split(",") if x.strip()]
        hypotheses = [x.strip() for x in data.get("hypotheses", "").split(",") if x.strip()]

        # Build reporting period from start and end dates
        start_date = data.get("start_date", "").strip()
        end_date = data.get("end_date", "").strip()
        reporting_period = f"{start_date} to {end_date}" if start_date and end_date else ""

        # Build briefing object
        briefing = {
            "brand": data.get("brand", "").strip(),
            "competitors": competitors,
            "market": data.get("market", "").strip(),
            "reporting_period": reporting_period,
            "objective": data.get("objective", "").strip(),
            "dashboard_links": dashboard_links,
            "research_urls": research_urls,
            "hypotheses": hypotheses
        }

        # Validate briefing
        validation_errors = validate_briefing(briefing)
        if validation_errors:
            return jsonify({"error": "Validation failed", "details": validation_errors}), 400

        # Ensure assistant exists
        assistant_id = ensure_assistant()

        # Create thread
        thread = client.beta.threads.create()
        print(f"Created thread: {thread.id}")

        # Initial message
        client.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content="Start new analysis. I will provide briefing and potentially a data file."
        )

        # Handle file upload if provided
        file_ids = []
        if "data_file" in request.files:
            file = request.files["data_file"]

            if file and file.filename:
                # Validate file
                if not allowed_file(file.filename):
                    return jsonify({
                        "error": f"Invalid file type. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
                    }), 400

                # Secure the filename
                filename = secure_filename(file.filename)

                # Save to temp file and upload to OpenAI
                with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{filename}") as tmp:
                    file.save(tmp.name)

                    try:
                        with open(tmp.name, "rb") as f:
                            uploaded_file = client.files.create(file=f, purpose="assistants")
                            file_ids.append(uploaded_file.id)
                            print(f"Uploaded file: {uploaded_file.id}")
                    finally:
                        # Clean up temp file
                        os.unlink(tmp.name)

                # Attach file to thread
                client.beta.threads.messages.create(
                    thread_id=thread.id,
                    role="user",
                    content=f"Here is the latest data export: {filename}",
                    attachments=[
                        {"file_id": fid, "tools": [{"type": "file_search"}, {"type": "code_interpreter"}]}
                        for fid in file_ids
                    ]
                )

        # Analyze dashboards with Playwright if provided
        dashboard_insights = {}
        if briefing['dashboard_links']:
            print(f"📊 Found {len(briefing['dashboard_links'])} dashboard link(s), scraping with Playwright...")
            dashboard_insights = analyze_dashboards_with_playwright(briefing['dashboard_links'])

        # Fetch competitor and research insights
        competitor_insights = ""
        all_research_urls = competitor_urls_list + research_urls
        if all_research_urls:
            print(f"🔍 Found {len(all_research_urls)} competitor/research URLs to analyze...")
            competitor_insights = fetch_competitor_insights(all_research_urls)

        # Send briefing
        briefing_content = f"""Here is the completed briefing form:

Brand: {briefing['brand']}
Market: {briefing['market']}
Reporting Period: {briefing['reporting_period']}
Objective: {briefing['objective']}
Competitors: {json.dumps(briefing['competitors'], indent=2)}
Dashboard Links: {', '.join(briefing['dashboard_links']) if briefing['dashboard_links'] else 'None'}
Research URLs: {', '.join(briefing['research_urls']) if briefing['research_urls'] else 'None'}
Hypotheses: {', '.join(briefing['hypotheses']) if briefing['hypotheses'] else 'None'}
"""
        client.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content=briefing_content
        )

        # Send dashboard insights if extracted
        if dashboard_insights:
            dashboard_content = "## Dashboard Insights Extracted by Manus AI\n\n"
            for url, insights in dashboard_insights.items():
                dashboard_content += f"### Dashboard: {url}\n\n{insights}\n\n---\n\n"

            client.beta.threads.messages.create(
                thread_id=thread.id,
                role="user",
                content=dashboard_content
            )
            print(f"✓ Sent dashboard insights to assistant ({len(dashboard_insights)} dashboard(s))")

        # Send competitor and research insights if fetched
        if competitor_insights:
            client.beta.threads.messages.create(
                thread_id=thread.id,
                role="user",
                content=competitor_insights
            )
            print(f"✓ Sent competitor/research insights to assistant ({len(all_research_urls)} URLs)")

        # Request analysis
        client.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content="Run full analysis using the core brief framework. Return output_schema JSON only."
        )

        # Create and run the assistant
        run = client.beta.threads.runs.create(
            thread_id=thread.id,
            assistant_id=assistant_id,
            instructions=(
                "CRITICAL: Analyze data from three sources: "
                "1) Uploaded files using file_search + code_interpreter, "
                "2) Scraped dashboard data provided in the briefing (tables, metrics, KPIs), "
                "3) Competitor and research insights extracted from provided URLs. "
                "Dashboard data is automatically extracted via browser automation and included in your context. "
                "Competitor/research insights are fetched and analyzed from the URLs provided in the briefing. "
                "If none of these data sources are available, state this clearly and do NOT fabricate metrics. "
                "NEVER make up numbers for impressions, clicks, CTR, conversions, or any performance data. "
                "Use the competitor insights to provide context, benchmarking, and competitive analysis. "
                "Respond strictly in output_schema JSON format with actual data from available sources."
            )
        )

        print(f"Started run: {run.id}")

        # Poll for completion
        final_run = poll_run_status(thread.id, run.id)
        print(f"Run completed: {final_run.id}")

        # Retrieve messages
        messages = client.beta.threads.messages.list(thread_id=thread.id)

        # Find the assistant's response
        for msg in messages.data:
            if msg.role == "assistant":
                # Extract text content
                parts = []
                for content_block in msg.content:
                    if content_block.type == "text":
                        parts.append(content_block.text.value)

                raw_response = "\n".join(parts).strip()

                # Try to parse as JSON
                try:
                    parsed_json = json.loads(raw_response)

                    # Log successful generation
                    duration = time.time() - start_time
                    log_usage('report_generation', {
                        'status': 'success',
                        'brand': data.get('brand', 'Unknown'),
                        'market': data.get('market', 'Unknown'),
                        'objective': data.get('objective', 'Unknown'),
                        'duration_seconds': round(duration, 2),
                        'duration_formatted': f"{int(duration // 60)}:{int(duration % 60):02d}",
                        'has_file_upload': 'data_file' in request.files,
                        'thread_id': thread.id,
                        'run_id': run.id
                    })

                    # Increment user's report counter if authenticated
                    if current_user.is_authenticated:
                        log_activity(
                            user_id=current_user.id,
                            action_type='report_generated',
                            details={
                                'brand': data.get('brand', 'Unknown'),
                                'market': data.get('market', 'Unknown'),
                                'duration_seconds': round(duration, 2)
                            },
                            resource_id=thread.id
                        )

                    return jsonify({
                        "success": True,
                        "data": parsed_json,
                        "thread_id": thread.id,
                        "run_id": run.id
                    })
                except json.JSONDecodeError:
                    # Return raw response if not valid JSON
                    return jsonify({
                        "success": False,
                        "error": "Response was not valid JSON",
                        "raw_response": raw_response,
                        "thread_id": thread.id
                    }), 500

        return jsonify({"error": "No assistant message found in thread"}), 500

    except TimeoutError as e:
        # Log failed generation
        duration = time.time() - start_time
        log_usage('report_generation', {
            'status': 'timeout',
            'brand': data.get('brand', 'Unknown') if 'data' in locals() else 'Unknown',
            'duration_seconds': round(duration, 2),
            'error': str(e)
        })
        return jsonify({"error": str(e)}), 408
    except Exception as e:
        # Log failed generation
        duration = time.time() - start_time
        log_usage('report_generation', {
            'status': 'error',
            'brand': data.get('brand', 'Unknown') if 'data' in locals() else 'Unknown',
            'duration_seconds': round(duration, 2),
            'error': str(e),
            'error_type': type(e).__name__
        })

        print(f"Error in /api/analyze: {str(e)}")
        if DEBUG:
            traceback.print_exc()
        return jsonify({
            "error": "Internal server error",
            "message": str(e),
            "type": type(e).__name__
        }), 500


@app.route("/api/elevenlabs-voices", methods=["GET"])
def get_elevenlabs_voices():
    """
    Fetch available voices from ElevenLabs API.
    """
    try:
        if not ELEVENLABS_API_KEY:
            return jsonify({
                "error": "ElevenLabs API key not configured",
                "message": "Please add ELEVENLABS_API_KEY to your .env file"
            }), 500

        # Call ElevenLabs API to get voices
        url = "https://api.elevenlabs.io/v1/voices"
        headers = {
            "xi-api-key": ELEVENLABS_API_KEY
        }

        response = requests.get(url, headers=headers)

        if response.status_code != 200:
            return jsonify({
                "error": "ElevenLabs API error",
                "message": response.text,
                "status_code": response.status_code
            }), response.status_code

        voices_data = response.json()

        # Format the response for easier frontend consumption
        formatted_voices = []
        for voice in voices_data.get("voices", []):
            formatted_voices.append({
                "voice_id": voice.get("voice_id"),
                "name": voice.get("name"),
                "category": voice.get("category", "unknown"),
                "labels": voice.get("labels", {}),
                "preview_url": voice.get("preview_url", "")
            })

        return jsonify({"voices": formatted_voices})

    except Exception as e:
        print(f"Error fetching ElevenLabs voices: {str(e)}")
        traceback.print_exc()
        return jsonify({
            "error": "Failed to fetch voices",
            "message": str(e)
        }), 500


# Quote of the Day caching
quote_cache = {
    'quote': 'The only way to do great work is to love what you do.',
    'author': 'Steve Jobs',
    'date': None
}


@app.route("/api/quote-of-the-day", methods=["GET"])
def get_quote_of_the_day():
    """
    Fetch the quote of the day from BrainyQuote.
    Caches the quote for 24 hours.
    """
    try:
        from datetime import date

        today = str(date.today())

        # Return cached quote if it's from today
        if quote_cache['date'] == today and quote_cache['quote']:
            return jsonify({
                'quote': quote_cache['quote'],
                'author': quote_cache['author'],
                'source': 'BrainyQuote',
                'cached': True
            })

        # Fetch fresh quote
        url = "https://www.brainyquote.com/quote_of_the_day"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate, br',
            'DNT': '1',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1'
        }

        response = requests.get(url, headers=headers, timeout=10)

        if response.status_code != 200:
            # Return cached quote (fallback)
            return jsonify({
                'quote': quote_cache['quote'],
                'author': quote_cache['author'],
                'source': 'BrainyQuote',
                'cached': True
            })

        soup = BeautifulSoup(response.content, 'html.parser')

        # Find the first quote of the day
        quote_elem = soup.select_one('.qotd-q-cntr .b-qt')
        author_elem = soup.select_one('.qotd-q-cntr .bq-aut')

        if not quote_elem:
            # Try alternative selectors
            quote_elem = soup.select_one('.oncl_q')
            author_elem = soup.select_one('.oncl_a')

        if quote_elem:
            quote_text = quote_elem.get_text(strip=True)
            author_text = author_elem.get_text(strip=True) if author_elem else "Unknown"

            # Update cache
            quote_cache['quote'] = quote_text
            quote_cache['author'] = author_text
            quote_cache['date'] = today

            return jsonify({
                'quote': quote_text,
                'author': author_text,
                'source': 'BrainyQuote',
                'cached': False
            })
        else:
            # Return cached quote (fallback)
            return jsonify({
                'quote': quote_cache['quote'],
                'author': quote_cache['author'],
                'source': 'BrainyQuote',
                'cached': True
            })

    except Exception as e:
        print(f"Error fetching quote: {str(e)}")
        traceback.print_exc()

        # Return cached quote (fallback)
        return jsonify({
            'quote': quote_cache['quote'],
            'author': quote_cache['author'],
            'source': 'BrainyQuote',
            'cached': True
        })


@app.route("/api/generate-audio", methods=["POST"])
def generate_audio():
    """
    Generate audio using ElevenLabs TTS API.
    Expects JSON: { "text": "script text", "voice_id": "voice_id_string" }
    """
    try:
        if not ELEVENLABS_API_KEY:
            return jsonify({
                "error": "ElevenLabs API key not configured",
                "message": "Please add ELEVENLABS_API_KEY to your .env file"
            }), 500

        data = request.get_json()
        text = data.get("text", "").strip()
        voice_id = data.get("voice_id", "").strip()

        if not text:
            return jsonify({"error": "No text provided"}), 400

        if not voice_id:
            return jsonify({"error": "No voice_id provided"}), 400

        # Call ElevenLabs API
        url = f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}"
        headers = {
            "Accept": "audio/mpeg",
            "Content-Type": "application/json",
            "xi-api-key": ELEVENLABS_API_KEY
        }
        payload = {
            "text": text,
            "model_id": "eleven_monolingual_v1",
            "voice_settings": {
                "stability": 0.5,
                "similarity_boost": 0.5
            }
        }

        response = requests.post(url, json=payload, headers=headers)

        if response.status_code != 200:
            return jsonify({
                "error": "ElevenLabs API error",
                "message": response.text,
                "status_code": response.status_code
            }), response.status_code

        # Return audio as binary response
        return Response(response.content, mimetype="audio/mpeg")

    except Exception as e:
        print(f"Error in generate_audio: {str(e)}")
        traceback.print_exc()
        return jsonify({
            "error": "Audio generation failed",
            "message": str(e)
        }), 500


def upload_file_to_topview(file_path, file_type='image'):
    """
    Upload a file to TopView AI using 3-step process:
    1. Get upload credentials
    2. Upload to S3
    3. Check upload status
    """
    headers = {
        "Topview-Uid": TOPVIEW_UID,
        "Authorization": f"Bearer {TOPVIEW_API_KEY}"
    }

    # Determine file format
    file_ext = os.path.splitext(file_path)[1].lower().lstrip('.')
    if file_type == 'image':
        # Supported: jpg, png, jpeg, bmp, webp
        format_map = {'jpg': 'jpg', 'jpeg': 'jpg', 'png': 'png', 'bmp': 'bmp', 'webp': 'webp'}
        file_format = format_map.get(file_ext, 'jpg')
    else:  # audio
        # Supported: mp3, wav, m4a
        format_map = {'mp3': 'mp3', 'wav': 'wav', 'm4a': 'm4a'}
        file_format = format_map.get(file_ext, 'mp3')

    print(f"Uploading {file_type} file: {file_path} (format: {file_format})")

    # Step 1: Get upload credentials
    print("Step 1: Getting upload credentials...")
    cred_response = requests.get(
        f"https://api.topview.ai/v1/upload/credential?format={file_format}",
        headers=headers,
        timeout=30
    )

    if cred_response.status_code != 200:
        raise Exception(f"Failed to get upload credentials: {cred_response.status_code} - {cred_response.text}")

    cred_data = cred_response.json()
    print(f"Credential response: {cred_data}")

    # Extract data from response
    result = cred_data.get('result', {})
    file_id = result.get('fileId')
    upload_url = result.get('uploadUrl')

    if not file_id or not upload_url:
        raise Exception(f"Missing fileId or uploadUrl in credential response: {cred_data}")

    print(f"Got fileId: {file_id}")
    print(f"Got uploadUrl: {upload_url[:100]}...")

    # Step 2: Upload file to S3
    print("Step 2: Uploading file to S3...")
    with open(file_path, 'rb') as f:
        file_data = f.read()

    upload_response = requests.put(
        upload_url,
        data=file_data,
        headers={'Content-Type': 'application/octet-stream'},
        timeout=120
    )

    if upload_response.status_code not in [200, 204]:
        raise Exception(f"S3 upload failed: {upload_response.status_code} - {upload_response.text[:200]}")

    print(f"S3 upload successful: {upload_response.status_code}")

    # Step 3: Check upload status
    print("Step 3: Checking upload status...")
    check_response = requests.get(
        f"https://api.topview.ai/v1/upload/check?fileId={file_id}",
        headers=headers,
        timeout=30
    )

    if check_response.status_code != 200:
        raise Exception(f"Upload check failed: {check_response.status_code} - {check_response.text}")

    check_data = check_response.json()
    print(f"Upload check response: {check_data}")

    # Verify the upload was successful
    if check_data.get('code') != '200':
        raise Exception(f"Upload verification failed: {check_data}")

    print(f"✓ Upload complete! File ID: {file_id}")
    return file_id


@app.route("/api/generate-lipsync", methods=["POST"])
def generate_lipsync():
    """
    Generate lipsync video using TopView AI Photo Avatar API.
    Expects multipart/form-data: image, audio, and optional prompt
    """
    start_time = time.time()
    video_url = None
    try:
        if not TOPVIEW_API_KEY or not TOPVIEW_UID:
            return jsonify({
                "error": "TopView AI credentials not configured",
                "message": "Please add TOPVIEW_API_KEY and TOPVIEW_UID to your .env file"
            }), 500

        # Get form data
        if 'image' not in request.files:
            return jsonify({"error": "No image file provided"}), 400

        if 'audio' not in request.files:
            return jsonify({"error": "No audio file provided"}), 400

        image_file = request.files['image']
        audio_file = request.files['audio']
        prompt = request.form.get('prompt', '')

        # Save files temporarily
        image_temp = tempfile.NamedTemporaryFile(delete=False, suffix='.jpg')
        audio_temp = tempfile.NamedTemporaryFile(delete=False, suffix='.mp3')

        image_file.save(image_temp.name)
        audio_file.save(audio_temp.name)

        # Get audio duration to estimate processing time
        print("Calculating audio duration...")
        try:
            duration_result = subprocess.run(
                ['ffprobe', '-v', 'error', '-show_entries', 'format=duration',
                 '-of', 'default=noprint_wrappers=1:nokey=1', audio_temp.name],
                capture_output=True,
                text=True,
                timeout=10
            )
            audio_duration_seconds = float(duration_result.stdout.strip())
            print(f"Audio duration: {audio_duration_seconds:.2f} seconds")

            # TopView takes ~45 seconds to render 1 second of video
            # Add 20% buffer for safety
            estimated_processing_time = int(audio_duration_seconds * 45 * 1.2)
            estimated_minutes = estimated_processing_time / 60
            print(f"Estimated processing time: {estimated_processing_time} seconds ({estimated_minutes:.1f} minutes)")
        except Exception as e:
            print(f"Warning: Could not calculate audio duration: {e}")
            audio_duration_seconds = 60  # Default to 60 seconds
            estimated_processing_time = 3600  # Default to 60 minutes
            estimated_minutes = 60

        print("Uploading image to TopView AI...")
        image_file_id = upload_file_to_topview(image_temp.name, 'image')
        print(f"Image uploaded, file ID: {image_file_id}")

        print("Uploading audio to TopView AI...")
        audio_file_id = upload_file_to_topview(audio_temp.name, 'audio')
        print(f"Audio uploaded, file ID: {audio_file_id}")

        # Clean up temp files
        os.unlink(image_temp.name)
        os.unlink(audio_temp.name)

        # Step 1: Submit task
        headers = {
            "Topview-Uid": TOPVIEW_UID,
            "Authorization": f"Bearer {TOPVIEW_API_KEY}",
            "Content-Type": "application/json"
        }

        payload = {
            "templateImageFileId": image_file_id,
            "mode": "avatar4",
            "scriptMode": "audio",
            "audioFileId": audio_file_id,
            "avatarActionPrompt": prompt if prompt else "",
            "saveCustomAiAvatar": False
        }

        print(f"Submitting task to TopView AI with payload: {payload}")

        response = requests.post(
            "https://api.topview.ai/v1/photo_avatar/task/submit",
            headers=headers,
            json=payload,
            timeout=30
        )

        if response.status_code != 200:
            return jsonify({
                "error": "TopView AI task submission failed",
                "message": f"Status {response.status_code}: {response.text}"
            }), response.status_code

        submit_result = response.json()
        print(f"Task submitted: {submit_result}")

        if submit_result.get('code') != '200':
            return jsonify({
                "error": "Task submission failed",
                "message": submit_result.get('message', 'Unknown error')
            }), 500

        task_id = submit_result.get('result', {}).get('taskId')
        if not task_id:
            return jsonify({
                "error": "No task ID returned",
                "message": str(submit_result)
            }), 500

        # Step 2: Poll for completion
        print(f"Polling for task completion: {task_id}")
        # Calculate max_attempts based on estimated processing time
        # Poll every 5 seconds, add 10 minutes buffer to estimated time
        max_attempts = int((estimated_processing_time + 600) / 5)  # +10 min buffer
        print(f"Max polling attempts: {max_attempts} (will wait up to {max_attempts * 5 / 60:.1f} minutes)")
        attempt = 0

        while attempt < max_attempts:
            time.sleep(5)  # Wait 5 seconds between polls
            attempt += 1

            query_response = requests.get(
                f"https://api.topview.ai/v1/photo_avatar/task/query?taskId={task_id}&needCloudFrontUrl",
                headers=headers,
                timeout=30
            )

            if query_response.status_code != 200:
                print(f"Query failed: {query_response.text}")
                continue

            query_result = query_response.json()
            print(f"Query attempt {attempt}: {query_result.get('result', {}).get('status')}")

            if query_result.get('code') != '200':
                continue

            result_data = query_result.get('result', {})
            status = result_data.get('status')

            if status == 'success':
                video_url = result_data.get('finishedVideoUrl')
                video_cover_url = result_data.get('finishedVideoCoverUrl')
                if video_url:
                    print(f"✓ Video generated successfully!")
                    print(f"  Video URL: {video_url[:100]}...")
                    print(f"  Cover URL: {video_cover_url[:100] if video_cover_url else 'N/A'}...")
                    print(f"  Full result: {result_data}")

                    # Download and re-encode video for better compatibility
                    print("Downloading video for re-encoding...")
                    try:
                        video_response = requests.get(video_url, timeout=120)
                        if video_response.status_code == 200:
                            # Save original video temporarily
                            original_video = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
                            original_video.write(video_response.content)
                            original_video.close()

                            print(f"Original video saved: {original_video.name}")

                            # Re-encode with FFmpeg for maximum compatibility
                            # H.264 codec with AAC audio is the most compatible format
                            reencoded_video = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
                            reencoded_path = reencoded_video.name
                            reencoded_video.close()

                            print("Re-encoding video with FFmpeg...")
                            ffmpeg_command = [
                                'ffmpeg',
                                '-i', original_video.name,
                                '-c:v', 'libx264',  # H.264 video codec
                                '-preset', 'fast',   # Encoding speed
                                '-crf', '23',        # Quality (lower = better, 23 is default)
                                '-c:a', 'aac',       # AAC audio codec
                                '-b:a', '128k',      # Audio bitrate
                                '-movflags', '+faststart',  # Enable progressive playback
                                '-pix_fmt', 'yuv420p',  # Pixel format for compatibility
                                '-y',  # Overwrite output file
                                reencoded_path
                            ]

                            print(f"FFmpeg command: {' '.join(ffmpeg_command)}")
                            result = subprocess.run(
                                ffmpeg_command,
                                capture_output=True,
                                text=True,
                                timeout=300  # 5 minutes timeout
                            )

                            print(f"FFmpeg return code: {result.returncode}")
                            if result.stdout:
                                print(f"FFmpeg stdout: {result.stdout[-500:]}")  # Last 500 chars
                            if result.stderr:
                                print(f"FFmpeg stderr: {result.stderr[-500:]}")  # Last 500 chars

                            if result.returncode == 0:
                                print(f"✓ Video re-encoded successfully: {reencoded_path}")
                                print(f"  Re-encoded file size: {os.path.getsize(reencoded_path)} bytes")

                                # Save to static folder for serving
                                static_dir = os.path.join(os.path.dirname(__file__), 'static', 'videos')
                                os.makedirs(static_dir, exist_ok=True)

                                video_filename = f"lipsync_{int(time.time())}.mp4"
                                final_path = os.path.join(static_dir, video_filename)

                                shutil.move(reencoded_path, final_path)

                                # Clean up original temp file
                                os.unlink(original_video.name)

                                # Return local URL instead of S3 URL
                                local_video_url = f"/static/videos/{video_filename}"
                                print(f"✓ Video saved to: {final_path}")
                                print(f"  Serving from: {local_video_url}")

                                # Log successful video generation
                                duration = time.time() - start_time
                                log_usage('video_generation', {
                                    'status': 'success',
                                    'duration_seconds': round(duration, 2),
                                    'duration_formatted': f"{int(duration // 60)}:{int(duration % 60):02d}",
                                    'video_filename': video_filename,
                                    'file_size_bytes': os.path.getsize(final_path),
                                    'file_size_kb': round(os.path.getsize(final_path) / 1024, 2),
                                    're_encoded': True,
                                    'prompt': prompt if prompt else 'None'
                                })

                                return jsonify({
                                    "success": True,
                                    "video_url": local_video_url,
                                    "cover_url": video_cover_url,
                                    "message": "Lipsync video generated and optimized successfully"
                                })
                            else:
                                print(f"FFmpeg error: {result.stderr}")
                                # Fall back to original URL if re-encoding fails
                                os.unlink(original_video.name)
                                if os.path.exists(reencoded_path):
                                    os.unlink(reencoded_path)

                    except Exception as reencode_error:
                        print(f"Re-encoding failed: {str(reencode_error)}")
                        traceback.print_exc()

                    # Fall back to original URL if anything goes wrong
                    duration = time.time() - start_time
                    log_usage('video_generation', {
                        'status': 'success',
                        'duration_seconds': round(duration, 2),
                        'duration_formatted': f"{int(duration // 60)}:{int(duration % 60):02d}",
                        're_encoded': False,
                        'fallback_to_original': True,
                        'prompt': prompt if 'prompt' in locals() else 'None'
                    })

                    return jsonify({
                        "success": True,
                        "video_url": video_url,
                        "cover_url": video_cover_url,
                        "message": "Lipsync video generated successfully"
                    })
                else:
                    print(f"ERROR: No video URL in success response: {result_data}")
            elif status == 'failed':
                error_msg = result_data.get('errorMsg', 'Task failed')

                # Log failed generation
                duration = time.time() - start_time
                log_usage('video_generation', {
                    'status': 'failed',
                    'duration_seconds': round(duration, 2),
                    'error': error_msg
                })

                return jsonify({
                    "error": "Video generation failed",
                    "message": error_msg
                }), 500

        # Log timeout
        duration = time.time() - start_time
        log_usage('video_generation', {
            'status': 'timeout',
            'duration_seconds': round(duration, 2),
            'estimated_minutes': estimated_minutes
        })

        return jsonify({
            "error": "Timeout",
            "message": f"Video generation timed out after {duration / 60:.1f} minutes. Expected time was {estimated_minutes:.1f} minutes. TopView may still be processing - check back later."
        }), 504

    except Exception as e:
        # Log error
        duration = time.time() - start_time
        log_usage('video_generation', {
            'status': 'error',
            'duration_seconds': round(duration, 2),
            'error': str(e),
            'error_type': type(e).__name__
        })

        print(f"Error in generate_lipsync: {str(e)}")
        traceback.print_exc()
        return jsonify({
            "error": "Lipsync generation failed",
            "message": str(e)
        }), 500


@app.route("/api/proxy-video")
def proxy_video():
    """
    Proxy video from external URL to avoid CORS issues.
    This allows videos from TopView AI's S3 to play in the browser.
    """
    video_url = request.args.get('url')

    if not video_url:
        return jsonify({"error": "No video URL provided"}), 400

    try:
        print(f"Proxying video from: {video_url}")

        # Fetch the video from the external URL
        response = requests.get(video_url, stream=True, timeout=30)

        if response.status_code != 200:
            return jsonify({
                "error": "Failed to fetch video",
                "status": response.status_code
            }), response.status_code

        # Get content type from the original response
        content_type = response.headers.get('Content-Type', 'video/mp4')

        # Stream the video content
        def generate():
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    yield chunk

        return Response(
            generate(),
            mimetype=content_type,
            headers={
                'Accept-Ranges': 'bytes',
                'Content-Type': content_type,
                'Access-Control-Allow-Origin': '*'
            }
        )

    except Exception as e:
        print(f"Error proxying video: {str(e)}")
        traceback.print_exc()
        return jsonify({
            "error": "Video proxy failed",
            "message": str(e)
        }), 500


@app.route("/static/videos/<filename>")
def serve_video(filename):
    """
    Serve video files with correct MIME type for browser compatibility.
    """
    try:
        videos_dir = os.path.join(os.path.dirname(__file__), 'static', 'videos')
        return send_from_directory(
            videos_dir,
            filename,
            mimetype='video/mp4',
            as_attachment=False,
            conditional=True
        )
    except Exception as e:
        print(f"Error serving video {filename}: {str(e)}")
        return jsonify({"error": "Video not found"}), 404


@app.route("/api/chat-modify", methods=["POST"])
def chat_modify():
    """
    Modify the report using OpenAI API.
    Expects JSON: { "message": "user instruction", "current_report": {...} }
    """
    try:
        data = request.get_json()
        message = data.get("message", "").strip()
        current_report = data.get("current_report", {})

        if not message:
            return jsonify({"error": "No message provided"}), 400

        if not current_report:
            return jsonify({"error": "No report data provided"}), 400

        # Build prompt for OpenAI
        system_prompt = """You are an AI assistant helping to modify marketing analysis reports.
You will receive a JSON report and a modification request from the user.
Your task is to modify the report according to the request while maintaining the exact JSON structure.

CRITICAL RULES:
1. Maintain the EXACT same JSON structure
2. Keep all section names: audience, media, creative, conversion, competitive, optimization, bonus, citations
3. Each section (except bonus and citations) must have: key_findings, supporting_data, research_context, implications, actions
4. Only modify content relevant to the user's request
5. Return ONLY valid JSON, no explanations"""

        user_prompt = f"""Current Report:
{json.dumps(current_report, indent=2)}

User's Modification Request: {message}

Please modify the report according to the user's request and return the complete modified report as valid JSON."""

        # Call OpenAI API with JSON mode
        response = client.chat.completions.create(
            model=MODEL_ID,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.7,
            max_tokens=4000
        )

        # Parse the response
        modified_report_text = response.choices[0].message.content
        print(f"OpenAI response length: {len(modified_report_text)}")

        try:
            modified_report = json.loads(modified_report_text)
            print(f"Successfully parsed JSON with keys: {list(modified_report.keys())}")

            return jsonify({
                "success": True,
                "report": modified_report,
                "message": "Report modified successfully"
            })

        except json.JSONDecodeError as e:
            print(f"JSON decode error: {str(e)}")
            return jsonify({
                "error": "Invalid JSON in response",
                "message": str(e),
                "original_response": modified_report_text[:500],
                "report": current_report
            }), 400

    except Exception as e:
        print(f"Error in chat_modify: {str(e)}")
        traceback.print_exc()
        return jsonify({
            "error": "Chat modification failed",
            "message": str(e)
        }), 500


@app.route("/api/generate-script", methods=["POST"])
def generate_script():
    """
    Generate a script from the report using OpenAI.
    Expects JSON: { "prompt": "user's script requirements", "report": {...} }
    """
    try:
        data = request.get_json()
        prompt = data.get("prompt", "").strip()
        report = data.get("report", {})

        if not prompt:
            return jsonify({"error": "No prompt provided"}), 400

        if not report:
            return jsonify({"error": "No report data provided"}), 400

        # Build prompt for OpenAI
        system_prompt = """You are a creative scriptwriter specializing in marketing and brand content.
Your task is to create engaging scripts from marketing analysis reports.
The scripts should be clear, concise, and suitable for video, audio, or presentation formats.
Focus on the most impactful insights and make them audience-friendly."""

        # Summarize the report for context
        report_summary = f"""Report Summary:
Sections covered: {', '.join(report.keys())}

Key Insights:
"""
        # Add top findings from each section
        for section_name in ['audience', 'media', 'creative', 'conversion', 'competitive', 'optimization']:
            if section_name in report and report[section_name]:
                section_data = report[section_name]
                if 'key_findings' in section_data and section_data['key_findings']:
                    report_summary += f"\n{section_name.title()}: {section_data['key_findings'][0]}"

        if 'bonus' in report and report['bonus']:
            bonus = report['bonus']
            if 'one_sentence' in bonus:
                report_summary += f"\n\nKey Takeaway: {bonus['one_sentence']}"

        user_prompt = f"""{report_summary}

User's Script Request: {prompt}

Please create a script based on the user's request. Make it engaging, clear, and focused on the most important insights from the report."""

        # Call OpenAI API
        response = client.chat.completions.create(
            model=MODEL_ID,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.8,
            max_tokens=1000
        )

        # Get the generated script
        script_text = response.choices[0].message.content.strip()
        print(f"Generated script length: {len(script_text)}")

        return jsonify({
            "success": True,
            "script": script_text,
            "message": "Script generated successfully"
        })

    except Exception as e:
        print(f"Error in generate_script: {str(e)}")
        traceback.print_exc()
        return jsonify({
            "error": "Script generation failed",
            "message": str(e)
        }), 500


@app.route("/api/export-txt", methods=["POST"])
def export_txt():
    """Export report as plain text file."""
    try:
        data = request.get_json()
        report = data.get("report", {})

        if not report:
            return jsonify({"error": "No report data provided"}), 400

        # Build text content
        text_lines = []
        text_lines.append("=" * 80)
        text_lines.append("SUPA REPORTS - ANALYSIS REPORT")
        text_lines.append("=" * 80)
        text_lines.append(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

        # Process each section
        sections = {
            "audience": "AUDIENCE & TARGETING INSIGHTS",
            "media": "MEDIA & CHANNEL EFFECTIVENESS",
            "creative": "CREATIVE PERFORMANCE & ENGAGEMENT",
            "conversion": "CONVERSION & PERFORMANCE DRIVERS",
            "competitive": "COMPETITIVE & MARKET INSIGHTS",
            "optimization": "OPTIMIZATION & NEXT STEPS"
        }

        for key, title in sections.items():
            if key in report and report[key]:
                section = report[key]
                text_lines.append(f"\n{'=' * 80}")
                text_lines.append(title)
                text_lines.append('=' * 80)

                for field_name, field_title in [
                    ("key_findings", "KEY FINDINGS"),
                    ("supporting_data", "SUPPORTING DATA"),
                    ("research_context", "RESEARCH CONTEXT"),
                    ("implications", "IMPLICATIONS"),
                    ("actions", "ACTIONS")
                ]:
                    if field_name in section and section[field_name]:
                        text_lines.append(f"\n{field_title}:")
                        for item in section[field_name]:
                            text_lines.append(f"  • {item}")

        # Add bonus section
        if "bonus" in report and report["bonus"]:
            bonus = report["bonus"]
            text_lines.append(f"\n{'=' * 80}")
            text_lines.append("BONUS INSIGHTS")
            text_lines.append('=' * 80)
            if "one_sentence" in bonus:
                text_lines.append(f"\nOne Sentence Summary:\n  {bonus['one_sentence']}")
            if "key_takeaway" in bonus:
                text_lines.append(f"\nKey Takeaway:\n  {bonus['key_takeaway']}")
            if "unexpected_learning" in bonus:
                text_lines.append(f"\nUnexpected Learning:\n  {bonus['unexpected_learning']}")

        text_content = "\n".join(text_lines)

        # Create file in memory
        buffer = BytesIO()
        buffer.write(text_content.encode('utf-8'))
        buffer.seek(0)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"supareports_analysis_{timestamp}.txt"

        return send_file(
            buffer,
            mimetype='text/plain',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        print(f"Error in export_txt: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": "Text export failed", "message": str(e)}), 500


@app.route("/api/export-pdf", methods=["POST"])
def export_pdf():
    """Export report as PDF file."""
    try:
        data = request.get_json()
        report = data.get("report", {})

        if not report:
            return jsonify({"error": "No report data provided"}), 400

        # Create PDF in memory
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        story = []

        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#000000',
            spaceAfter=12,
            alignment=1  # Center
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='#50C878',
            spaceAfter=8
        )
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontSize=11,
            textColor='#000000',
            spaceAfter=6
        )
        body_style = styles['BodyText']

        # Add title
        story.append(Paragraph("SUPA REPORTS", title_style))
        story.append(Paragraph(f"Analysis Report - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))

        # Process each section
        sections = {
            "audience": "Audience & Targeting Insights",
            "media": "Media & Channel Effectiveness",
            "creative": "Creative Performance & Engagement",
            "conversion": "Conversion & Performance Drivers",
            "competitive": "Competitive & Market Insights",
            "optimization": "Optimization & Next Steps"
        }

        for key, title in sections.items():
            if key in report and report[key]:
                section = report[key]
                story.append(Paragraph(title, heading_style))
                story.append(Spacer(1, 0.1*inch))

                for field_name, field_title in [
                    ("key_findings", "Key Findings"),
                    ("supporting_data", "Supporting Data"),
                    ("research_context", "Research Context"),
                    ("implications", "Implications"),
                    ("actions", "Actions")
                ]:
                    if field_name in section and section[field_name]:
                        story.append(Paragraph(field_title, subheading_style))
                        for item in section[field_name]:
                            story.append(Paragraph(f"• {item}", body_style))
                        story.append(Spacer(1, 0.1*inch))

                story.append(Spacer(1, 0.2*inch))

        # Add bonus section
        if "bonus" in report and report["bonus"]:
            bonus = report["bonus"]
            story.append(Paragraph("Bonus Insights", heading_style))
            story.append(Spacer(1, 0.1*inch))
            if "one_sentence" in bonus:
                story.append(Paragraph("<b>One Sentence Summary:</b>", subheading_style))
                story.append(Paragraph(bonus['one_sentence'], body_style))
                story.append(Spacer(1, 0.1*inch))
            if "key_takeaway" in bonus:
                story.append(Paragraph("<b>Key Takeaway:</b>", subheading_style))
                story.append(Paragraph(bonus['key_takeaway'], body_style))
                story.append(Spacer(1, 0.1*inch))
            if "unexpected_learning" in bonus:
                story.append(Paragraph("<b>Unexpected Learning:</b>", subheading_style))
                story.append(Paragraph(bonus['unexpected_learning'], body_style))

        # Build PDF
        doc.build(story)
        buffer.seek(0)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"supareports_analysis_{timestamp}.pdf"

        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        print(f"Error in export_pdf: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": "PDF export failed", "message": str(e)}), 500


@app.route("/api/export-docx", methods=["POST"])
def export_docx():
    """Export report as Word document."""
    try:
        data = request.get_json()
        report = data.get("report", {})

        if not report:
            return jsonify({"error": "No report data provided"}), 400

        # Create Word document
        doc = Document()

        # Add title
        title = doc.add_heading('SUPA REPORTS', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        subtitle = doc.add_paragraph(f"Analysis Report - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Process each section
        sections = {
            "audience": "Audience & Targeting Insights",
            "media": "Media & Channel Effectiveness",
            "creative": "Creative Performance & Engagement",
            "conversion": "Conversion & Performance Drivers",
            "competitive": "Competitive & Market Insights",
            "optimization": "Optimization & Next Steps"
        }

        for key, title_text in sections.items():
            if key in report and report[key]:
                section = report[key]

                # Add section heading
                heading = doc.add_heading(title_text, 1)
                run = heading.runs[0]
                run.font.color.rgb = RGBColor(80, 200, 120)  # Emerald color

                for field_name, field_title in [
                    ("key_findings", "Key Findings"),
                    ("supporting_data", "Supporting Data"),
                    ("research_context", "Research Context"),
                    ("implications", "Implications"),
                    ("actions", "Actions")
                ]:
                    if field_name in section and section[field_name]:
                        doc.add_heading(field_title, 2)
                        for item in section[field_name]:
                            p = doc.add_paragraph(item, style='List Bullet')

        # Add bonus section
        if "bonus" in report and report["bonus"]:
            bonus = report["bonus"]
            heading = doc.add_heading("Bonus Insights", 1)
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(80, 200, 120)

            if "one_sentence" in bonus:
                doc.add_heading("One Sentence Summary", 2)
                doc.add_paragraph(bonus['one_sentence'])
            if "key_takeaway" in bonus:
                doc.add_heading("Key Takeaway", 2)
                doc.add_paragraph(bonus['key_takeaway'])
            if "unexpected_learning" in bonus:
                doc.add_heading("Unexpected Learning", 2)
                doc.add_paragraph(bonus['unexpected_learning'])

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"supareports_analysis_{timestamp}.docx"

        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        print(f"Error in export_docx: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": "Word export failed", "message": str(e)}), 500


@app.errorhandler(413)
def request_entity_too_large(error):
    """Handle file too large errors."""
    return jsonify({
        "error": "File too large",
        "max_size_mb": MAX_FILE_SIZE / (1024 * 1024)
    }), 413


@app.errorhandler(500)
def internal_server_error(error):
    """Handle internal server errors."""
    return jsonify({"error": "Internal server error"}), 500


# ============================================================================
# Main
# ============================================================================

if __name__ == "__main__":
    # Validate configuration
    if not OPENAI_API_KEY or OPENAI_API_KEY == "your-api-key-here":
        print("ERROR: Please set OPENAI_API_KEY in your .env file")
        exit(1)

    print("=" * 70)
    print("SupaReports — Upgrade Your Insights")
    print("=" * 70)
    print(f"Model: {MODEL_ID}")
    print(f"Assistant ID: {ASSISTANT_ID or 'Will create new'}")
    print(f"Debug mode: {DEBUG}")
    print(f"Max file size: {MAX_FILE_SIZE / (1024 * 1024):.1f}MB")
    print("=" * 70)

    # Use waitress for production-like server on macOS
    try:
        from waitress import serve
        print(f"\n🚀 Serving on http://localhost:{PORT}")
        print(f"Press Ctrl+C to stop\n")
        serve(app, host="0.0.0.0", port=PORT)
    except ImportError:
        print("\nWARNING: waitress not installed, falling back to Flask dev server")
        print(f"🚀 Serving on http://localhost:{PORT}")
        print(f"Press Ctrl+C to stop\n")
        app.run(host="0.0.0.0", port=PORT, debug=DEBUG)
